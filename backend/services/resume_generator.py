from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from io import BytesIO
import json

def generate_resume_pdf(resume_data: dict) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=LETTER)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, spaceBefore=12, spaceAfter=6)
    body_style = styles['Normal']
    
    elements = []
    
    # Contact Info
    elements.append(Paragraph(resume_data.get('contact_info', 'Resume'), title_style))
    
    # Social Links
    social_text = []
    for link in resume_data.get('social_links', []):
        label = link.get('label', 'Link')
        url = link.get('url', '#')
        social_text.append(f'<a href="{url}" color="blue">{label}</a>')
    
    if social_text:
        elements.append(Paragraph(" | ".join(social_text), body_style))
        
    elements.append(Spacer(1, 12))
    
    # Summary
    elements.append(Paragraph("Professional Summary", heading_style))
    elements.append(Paragraph(resume_data.get('summary', ''), body_style))
    
    # Experience
    elements.append(Paragraph("Experience", heading_style))
    for exp in resume_data.get('experience', []):
        exp_title = f"<b>{exp.get('role', '')}</b> at {exp.get('company', '')} ({exp.get('duration', '')})"
        elements.append(Paragraph(exp_title, body_style))
        for bullet in exp.get('bullets', []):
            elements.append(Paragraph(f"• {bullet}", body_style))
        elements.append(Spacer(1, 6))
        
    # Skills
    elements.append(Paragraph("Skills", heading_style))
    elements.append(Paragraph(", ".join(resume_data.get('skills', [])), body_style))
    
    # Education
    elements.append(Paragraph("Education", heading_style))
    for edu in resume_data.get('education', []):
        edu_info = f"<b>{edu.get('degree', '')}</b>, {edu.get('institution', '')} ({edu.get('duration', '')})"
        elements.append(Paragraph(edu_info, body_style))
        
    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_resume_docx(resume_data: dict) -> BytesIO:
    buffer = BytesIO()
    doc = Document()
    
    # Title / Contact
    doc.add_heading(resume_data.get('contact_info', 'Resume'), 0)
    
    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(resume_data.get('summary', ''))
    
    # Experience
    doc.add_heading('Experience', level=1)
    for exp in resume_data.get('experience', []):
        p = doc.add_paragraph()
        p.add_run(f"{exp.get('role', '')} at {exp.get('company', '')}").bold = True
        p.add_run(f"\n{exp.get('duration', '')}").italic = True
        
        for bullet in exp.get('bullets', []):
            doc.add_paragraph(bullet, style='List Bullet')
            
    # Skills
    doc.add_heading('Skills', level=1)
    doc.add_paragraph(", ".join(resume_data.get('skills', [])))
    
    # Education
    doc.add_heading('Education', level=1)
    for edu in resume_data.get('education', []):
        doc.add_paragraph(f"{edu.get('degree', '')}, {edu.get('institution', '')} ({edu.get('duration', '')})")
        
    doc.save(buffer)
    buffer.seek(0)
    return buffer
